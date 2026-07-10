import re
from html import escape
from io import BytesIO
from typing import Any

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


SCORING_DIMENSIONS = (
    ("skills_match", "Skills Match"),
    ("project_experience", "Project Experience"),
    ("education", "Education"),
    ("work_experience", "Work Experience"),
    ("keyword_match", "Keyword Match"),
)

ATS_FIELDS = (
    ("important_keywords", "Important Keywords"),
    ("matched_keywords", "Matched Keywords"),
    ("missing_keywords", "Missing Keywords"),
    ("keyword_suggestions", "Keyword Suggestions"),
)


def clean_text(value: Any, fallback: str = "") -> str:
    text = str(value or "").strip()
    return text or fallback


def safe_float(value: Any) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def safe_int(value: Any) -> int | None:
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def format_duration(
    duration_ms: Any,
    duration_us: Any = None,
    *,
    status: Any = "",
    completed_zero_as_sub_ms: bool = True,
) -> str:
    clean_status = clean_text(status).lower()
    if clean_status == "skipped":
        return "Skipped"

    parsed_us = safe_int(duration_us)
    if parsed_us is not None and 0 < parsed_us < 1000:
        return "<1 ms"

    parsed_ms = safe_float(duration_ms)
    if parsed_ms is None:
        return "N/A"
    if 0 < parsed_ms < 1:
        return "<1 ms"
    if parsed_ms == 0 and completed_zero_as_sub_ms and clean_status in {"completed", "failed"}:
        return "<1 ms"
    if parsed_ms < 1000:
        return f"{parsed_ms:.2f} ms"
    return f"{parsed_ms / 1000:.2f} s"


def as_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item or "").strip()]


def as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def safe_filename_part(value: Any, fallback: str) -> str:
    text = clean_text(value, fallback)
    text = re.sub(r"[^A-Za-z0-9]+", "-", text).strip("-").lower()
    return text[:60] or fallback


def build_export_filename(prefix: str, record: dict[str, Any], extension: str) -> str:
    company = safe_filename_part(record.get("company_name"), "unknown-company")
    title = safe_filename_part(record.get("job_title"), "unknown-position")
    return f"{prefix}-{company}-{title}.{extension}"


def build_cover_letter_docx(record: dict[str, Any]) -> BytesIO:
    document = Document()
    document.add_heading("Cover Letter", level=1)
    document.add_paragraph(f"Company Name: {clean_text(record.get('company_name'), 'Unknown Company')}")
    document.add_paragraph(f"Job Title: {clean_text(record.get('job_title'), 'Unknown Position')}")
    document.add_heading("Generated Cover Letter", level=2)

    cover_letter = clean_text(record.get("cover_letter"), "No cover letter generated.")
    for paragraph in cover_letter.splitlines() or [cover_letter]:
        document.add_paragraph(paragraph)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def get_pdf_font_name() -> str:
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        return "Helvetica"


def paragraph_text(value: Any, fallback: str = "Not provided") -> str:
    text = clean_text(value, fallback)
    return escape(text).replace("\n", "<br/>")


def add_heading(story: list[Any], text: str, style: ParagraphStyle) -> None:
    story.append(Paragraph(escape(text), style))
    story.append(Spacer(1, 0.08 * inch))


def add_paragraph_block(
    story: list[Any],
    title: str,
    value: Any,
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
    fallback: str = "Not provided",
) -> None:
    add_heading(story, title, heading_style)
    story.append(Paragraph(paragraph_text(value, fallback), body_style))
    story.append(Spacer(1, 0.14 * inch))


def add_list_block(
    story: list[Any],
    title: str,
    items: Any,
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
) -> None:
    add_heading(story, title, heading_style)
    safe_items = as_list(items)
    if not safe_items:
        story.append(Paragraph("No items found.", body_style))
    else:
        for item in safe_items:
            story.append(Paragraph(f"- {paragraph_text(item)}", body_style))
    story.append(Spacer(1, 0.14 * inch))


def add_summary_table(
    story: list[Any],
    record: dict[str, Any],
    body_style: ParagraphStyle,
) -> None:
    data = [
        ["Company Name", paragraph_text(record.get("company_name"), "Unknown Company")],
        ["Job Title", paragraph_text(record.get("job_title"), "Unknown Position")],
        ["Job URL", paragraph_text(record.get("job_url"))],
        ["Application Status", paragraph_text(record.get("application_status"), "Saved")],
        ["Match Score", f"{clean_text(record.get('match_score'), '0')}/100"],
        ["RAG Mode", paragraph_text(record.get("rag_mode"), "Not recorded")],
        ["Security Status", paragraph_text(record.get("security_status"), "not_available")],
    ]
    table = Table(
        [
            [
                Paragraph(f"<b>{escape(str(label))}</b>", body_style),
                Paragraph(str(value), body_style),
            ]
            for label, value in data
        ],
        colWidths=[1.75 * inch, 4.75 * inch],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#edf2f7")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.18 * inch))


def add_scoring_breakdown(
    story: list[Any],
    record: dict[str, Any],
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
) -> None:
    add_heading(story, "Scoring Breakdown", heading_style)
    breakdown = as_dict(record.get("scoring_breakdown"))

    for key, label in SCORING_DIMENSIONS:
        section = as_dict(breakdown.get(key))
        score = clean_text(section.get("score"), "0")
        reason = paragraph_text(section.get("reason"), "No reason generated.")
        story.append(Paragraph(f"<b>{escape(label)}: {escape(score)}/100</b>", body_style))
        story.append(Paragraph(reason, body_style))
        evidence = as_list(section.get("evidence"))
        if evidence:
            for item in evidence:
                story.append(Paragraph(f"- {paragraph_text(item)}", body_style))
        else:
            story.append(Paragraph("- No evidence provided.", body_style))
        story.append(Spacer(1, 0.08 * inch))

    story.append(Spacer(1, 0.06 * inch))


def add_ats_analysis(
    story: list[Any],
    record: dict[str, Any],
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
) -> None:
    add_heading(story, "ATS Keyword Analysis", heading_style)
    ats_analysis = as_dict(record.get("ats_analysis"))
    for key, label in ATS_FIELDS:
        story.append(Paragraph(f"<b>{escape(label)}</b>", body_style))
        items = as_list(ats_analysis.get(key))
        if items:
            story.append(Paragraph(paragraph_text(", ".join(items)), body_style))
        else:
            story.append(Paragraph("No keywords found.", body_style))
        story.append(Spacer(1, 0.08 * inch))

    story.append(Spacer(1, 0.06 * inch))


def add_upgraded_bullets(
    story: list[Any],
    record: dict[str, Any],
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
) -> None:
    add_heading(story, "Upgraded Resume Bullets", heading_style)
    bullets = record.get("upgraded_resume_bullets")
    if not isinstance(bullets, list) or not bullets:
        story.append(Paragraph("No bullet improvements generated.", body_style))
        story.append(Spacer(1, 0.14 * inch))
        return

    for item in bullets:
        bullet = as_dict(item)
        original = paragraph_text(bullet.get("original"), "Not provided")
        improved = paragraph_text(bullet.get("improved"), "Not provided")
        reason = paragraph_text(bullet.get("reason"), "Not provided")
        story.append(Paragraph(f"<b>Original:</b> {original}", body_style))
        story.append(Paragraph(f"<b>Improved:</b> {improved}", body_style))
        story.append(Paragraph(f"<b>Reason:</b> {reason}", body_style))
        story.append(Spacer(1, 0.1 * inch))

    story.append(Spacer(1, 0.06 * inch))


def add_rag_sources(
    story: list[Any],
    record: dict[str, Any],
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
) -> None:
    add_heading(story, "RAG Sources", heading_style)
    story.append(
        Paragraph(f"RAG Mode: {paragraph_text(record.get('rag_mode'), 'Not recorded')}", body_style)
    )
    story.append(Spacer(1, 0.06 * inch))

    sources = record.get("rag_sources")
    if not isinstance(sources, list) or not sources:
        story.append(Paragraph("No RAG sources used.", body_style))
        story.append(Spacer(1, 0.14 * inch))
        return

    for source in sources:
        item = as_dict(source)
        title = paragraph_text(item.get("document_title"), "Untitled knowledge document")
        category = paragraph_text(item.get("category"), "Other")
        chunk_index = paragraph_text(item.get("chunk_index"), "0")
        reason = paragraph_text(item.get("relevance_reason"), "No relevance reason provided.")
        story.append(Paragraph(f"<b>{title}</b>", body_style))
        story.append(Paragraph(f"Category: {category} | Chunk Index: {chunk_index}", body_style))
        story.append(Paragraph(f"Reason: {reason}", body_style))
        story.append(Spacer(1, 0.1 * inch))

    story.append(Spacer(1, 0.06 * inch))


def add_agent_workflow(
    story: list[Any],
    record: dict[str, Any],
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
) -> None:
    add_heading(story, "Agent Workflow", heading_style)
    steps = record.get("workflow_steps")
    if not isinstance(steps, list) or not steps:
        story.append(
            Paragraph("No workflow audit trail is available for this older record.", body_style)
        )
        story.append(Spacer(1, 0.14 * inch))
        return

    workflow_duration = record.get("workflow_duration_ms")
    if workflow_duration not in (None, ""):
        duration_text = format_duration(
            workflow_duration,
            record.get("workflow_duration_us"),
            completed_zero_as_sub_ms=False,
        )
        story.append(Paragraph(f"Total workflow duration: {escape(duration_text)}", body_style))
        story.append(Spacer(1, 0.06 * inch))

    for step in steps:
        item = as_dict(step)
        name = paragraph_text(item.get("name"), "Unnamed Step")
        status = paragraph_text(item.get("status"), "pending")
        message = paragraph_text(item.get("message"), "No message recorded.")
        duration = format_duration(
            item.get("duration_ms"),
            item.get("duration_us"),
            status=item.get("status"),
        )
        story.append(Paragraph(f"<b>{name}</b> - {status} - {escape(duration)}", body_style))
        story.append(Paragraph(message, body_style))
        story.append(Spacer(1, 0.06 * inch))

    story.append(Spacer(1, 0.06 * inch))


def add_security_audit(
    story: list[Any],
    record: dict[str, Any],
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
) -> None:
    add_heading(story, "AI Security Audit", heading_style)
    scan = as_dict(record.get("security_scan"))
    summary = as_dict(scan.get("redaction_summary"))
    status = clean_text(record.get("security_status"), "not_available")
    policy_version = clean_text(record.get("security_policy_version"), "Not recorded")
    risk_level = clean_text(scan.get("risk_level"), "not_available")
    prompt_injection = "Yes" if scan.get("prompt_injection_detected") else "No"
    sensitive_data = "Yes" if scan.get("sensitive_data_detected") else "No"
    redaction_counts = (
        f"Email: {clean_text(summary.get('email_count'), '0')}; "
        f"Phone: {clean_text(summary.get('phone_count'), '0')}; "
        f"Secrets: {clean_text(summary.get('secret_count'), '0')}; "
        f"Private keys: {clean_text(summary.get('private_key_count'), '0')}"
    )

    for label, value in (
        ("Policy Version", policy_version),
        ("Security Status", status),
        ("Risk Level", risk_level),
        ("Prompt Injection Detected", prompt_injection),
        ("Sensitive Credential Detected", sensitive_data),
        ("PII Redaction Counts", redaction_counts),
    ):
        story.append(Paragraph(f"<b>{escape(label)}:</b> {paragraph_text(value)}", body_style))

    findings = scan.get("findings")
    story.append(Paragraph("<b>Safe Findings Summary</b>", body_style))
    if isinstance(findings, list) and findings:
        for item in findings:
            finding = as_dict(item)
            category = paragraph_text(finding.get("category"), "security")
            severity = paragraph_text(finding.get("severity"), "info")
            source = paragraph_text(finding.get("source"), "unknown")
            message = paragraph_text(finding.get("message"), "Security finding detected.")
            story.append(
                Paragraph(
                    f"- {category} | {severity} | {source}: {message}",
                    body_style,
                )
            )
    else:
        story.append(Paragraph("No security findings recorded.", body_style))

    story.append(Spacer(1, 0.14 * inch))


def add_next_action(
    story: list[Any],
    record: dict[str, Any],
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
) -> None:
    add_heading(story, "Recommended Next Action", heading_style)
    next_action = as_dict(record.get("next_action"))
    if not next_action or not clean_text(next_action.get("action")):
        story.append(Paragraph("No next-action recommendation is available.", body_style))
        story.append(Spacer(1, 0.14 * inch))
        return

    label = paragraph_text(next_action.get("label"), "No Recommendation")
    priority = paragraph_text(next_action.get("priority"), "low")
    confidence = clean_text(next_action.get("confidence"), "0")
    reason = paragraph_text(next_action.get("reason"), "No reason recorded.")
    story.append(Paragraph(f"<b>{label}</b>", body_style))
    story.append(Paragraph(f"Priority: {priority}", body_style))
    story.append(Paragraph(f"Rule-based confidence: {escape(confidence)}", body_style))
    story.append(Paragraph(f"Reason: {reason}", body_style))

    tasks = as_list(next_action.get("recommended_tasks"))
    story.append(Paragraph("<b>Recommended Tasks</b>", body_style))
    if tasks:
        for task in tasks:
            story.append(Paragraph(f"- {paragraph_text(task)}", body_style))
    else:
        story.append(Paragraph("No tasks recorded.", body_style))

    evidence = as_list(next_action.get("evidence"))
    story.append(Paragraph("<b>Evidence</b>", body_style))
    if evidence:
        for item in evidence:
            story.append(Paragraph(f"- {paragraph_text(item)}", body_style))
    else:
        story.append(Paragraph("No evidence recorded.", body_style))

    story.append(Spacer(1, 0.14 * inch))


def add_human_decision(
    story: list[Any],
    record: dict[str, Any],
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
) -> None:
    add_heading(story, "Human Decision", heading_style)
    story.append(
        Paragraph(
            f"Decision: {paragraph_text(record.get('next_action_decision'), 'pending')}",
            body_style,
        )
    )
    story.append(
        Paragraph(
            f"Notes: {paragraph_text(record.get('next_action_decision_notes'), 'No notes recorded.')}",
            body_style,
        )
    )
    story.append(
        Paragraph(
            f"Decided at: {paragraph_text(record.get('next_action_decided_at'), 'Not decided')}",
            body_style,
        )
    )
    story.append(Spacer(1, 0.14 * inch))


def build_analysis_report_pdf(record: dict[str, Any]) -> BytesIO:
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        title="Personal Job Application Analysis Report",
    )

    font_name = get_pdf_font_name()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ExportTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#12332d"),
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "ExportHeading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#166b5f"),
        spaceBefore=6,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "ExportBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=9.5,
        leading=13,
        spaceAfter=3,
    )

    story: list[Any] = [
        Paragraph("Personal Job Application Analysis Report", title_style),
    ]
    add_summary_table(story, record, body_style)
    add_paragraph_block(story, "Match Reason", record.get("match_reason"), heading_style, body_style)
    add_paragraph_block(story, "Job Summary", record.get("job_summary"), heading_style, body_style)
    add_scoring_breakdown(story, record, heading_style, body_style)
    add_ats_analysis(story, record, heading_style, body_style)
    add_list_block(story, "Matched Skills", record.get("matched_skills"), heading_style, body_style)
    add_list_block(story, "Missing Skills", record.get("missing_skills"), heading_style, body_style)
    add_list_block(
        story,
        "Resume Suggestions",
        record.get("resume_suggestions"),
        heading_style,
        body_style,
    )
    add_upgraded_bullets(story, record, heading_style, body_style)
    add_rag_sources(story, record, heading_style, body_style)
    add_security_audit(story, record, heading_style, body_style)
    add_agent_workflow(story, record, heading_style, body_style)
    add_next_action(story, record, heading_style, body_style)
    add_human_decision(story, record, heading_style, body_style)
    add_paragraph_block(
        story,
        "Cover Letter",
        record.get("cover_letter"),
        heading_style,
        body_style,
        fallback="No cover letter generated.",
    )

    document.build(story)
    buffer.seek(0)
    return buffer
