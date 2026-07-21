"""Resume library, safe import, immutable versions, and structured diff."""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from typing import Any
from uuid import UUID

import chardet
from docx import Document
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import V2Settings
from app.db.models import FileAsset, Resume, ResumeVersion, utc_now
from app.db.repositories.auth import AuthRepository
from app.resumes.repository import ResumeRepository
from app.resumes.schemas import ResumeContent
from app.storage.local import LocalStorageProvider
from app.storage.validation import (
    DOCX_MEDIA_TYPE,
    MARKDOWN_MEDIA_TYPE,
    PDF_MEDIA_TYPE,
    TEXT_MEDIA_TYPE,
    validate_resume_upload,
)


class ResumeNotFound(RuntimeError):
    pass


class ResumeConflict(RuntimeError):
    pass


def _serialize(value: object, *, include_content: bool = False) -> dict[str, object]:
    if isinstance(value, Resume):
        return {
            "id": str(value.id),
            "title": value.title,
            "language": value.language,
            "target_role": value.target_role,
            "status": value.status,
            "is_primary": value.is_primary,
            "active_version_id": str(value.active_version_id) if value.active_version_id else None,
            "created_at": value.created_at,
            "updated_at": value.updated_at,
            "archived_at": value.archived_at,
        }
    if isinstance(value, ResumeVersion):
        result = {
            "id": str(value.id),
            "resume_id": str(value.resume_id),
            "version_number": value.version_number,
            "parent_version_id": str(value.parent_version_id) if value.parent_version_id else None,
            "source_type": value.source_type,
            "source_file_id": str(value.source_file_id) if value.source_file_id else None,
            "schema_version": value.schema_version,
            "change_summary": value.change_summary,
            "status": value.status,
            "created_at": value.created_at,
            "finalized_at": value.finalized_at,
        }
        if include_content:
            result["content"] = value.content_json
            result["parsed_text"] = value.parsed_text
            result["extracted_text"] = value.parsed_text
        return result
    if isinstance(value, FileAsset):
        return {
            "id": str(value.id),
            "kind": value.kind,
            "original_filename": value.original_filename,
            "media_type": value.media_type,
            "mime_type": value.media_type,
            "size_bytes": value.size_bytes,
            "file_size": value.size_bytes,
            "sha256": value.sha256,
            "content_hash": value.sha256,
            "created_at": value.created_at,
            "deleted_at": value.deleted_at,
        }
    raise TypeError("Unsupported serialization type.")


class ResumeService:
    def __init__(self, db: Session, user_id: UUID, settings: V2Settings):
        self.db = db
        self.user_id = user_id
        self.settings = settings
        self.repository = ResumeRepository(db)
        self.storage = LocalStorageProvider(settings.file_storage_root)

    def list(self) -> list[dict[str, object]]:
        return [_serialize(value) for value in self.repository.resumes(self.user_id)]

    def primary(self) -> dict[str, object] | None:
        value = self.repository.primary(self.user_id)
        if value is None:
            return None
        result = _serialize(value)
        active_version = (
            self.repository.version(value.id, value.active_version_id)
            if value.active_version_id is not None
            else None
        )
        result["active_version"] = (
            _serialize(active_version, include_content=True) if active_version is not None else None
        )
        return result

    def create(self, values: dict[str, object]) -> dict[str, object]:
        resume = Resume(user_id=self.user_id, **values)
        self.db.add(resume)
        self.db.flush()
        if self.repository.primary(self.user_id) is None:
            self._make_primary(resume)
        self._audit("resume.created", resume.id)
        return _serialize(resume)

    def get(self, resume_id: UUID) -> dict[str, object]:
        return _serialize(self._resume(resume_id))

    def update(self, resume_id: UUID, values: dict[str, object]) -> dict[str, object]:
        resume = self._resume_for_update(resume_id)
        for key, value in values.items():
            if value is not None:
                setattr(resume, key, value)
        self._audit("resume.updated", resume.id)
        return _serialize(resume)

    def archive(self, resume_id: UUID) -> None:
        resume = self._resume_for_update(resume_id)
        was_primary = resume.is_primary
        resume.is_primary = False
        resume.status = "archived"
        resume.archived_at = utc_now()
        if was_primary:
            remaining = [item for item in self.repository.active_for_update(self.user_id) if item.id != resume.id]
            if remaining:
                self._make_primary(remaining[0], locked=remaining)
        self._audit("resume.archived", resume.id)

    def versions(self, resume_id: UUID) -> list[dict[str, object]]:
        resume = self._resume(resume_id)
        return [_serialize(value) for value in self.repository.versions(resume.id)]

    def version(self, resume_id: UUID, version_id: UUID, include_content: bool = True) -> dict[str, object]:
        resume = self._resume(resume_id)
        return _serialize(self._version(resume.id, version_id), include_content=include_content)

    def analysis_text(self, version_id: UUID) -> str:
        version = self.repository.owned_version(self.user_id, version_id)
        if version is None:
            raise ResumeNotFound("Resume Version not found.")
        if version.parsed_text.strip():
            return version.parsed_text
        values: list[str] = []

        def collect(value: object) -> None:
            if isinstance(value, str) and value.strip():
                values.append(value.strip())
            elif isinstance(value, dict):
                for nested in value.values():
                    collect(nested)
            elif isinstance(value, list):
                for nested in value:
                    collect(nested)

        collect(version.content_json)
        text = "\n".join(values)
        if not text:
            raise ResumeConflict("Resume Version has no analyzable content.")
        return text

    def create_version(
        self,
        resume_id: UUID,
        content: dict[str, object],
        parent_version_id: UUID | None,
        change_summary: str,
        parsed_text: str = "",
        source_type: str = "manual",
        source_file_id: UUID | None = None,
    ) -> dict[str, object]:
        resume = self._resume_for_update(resume_id)
        parent = None
        if parent_version_id:
            parent = self._version(resume.id, parent_version_id)
        validated = ResumeContent.model_validate(content)
        version = ResumeVersion(
            resume_id=resume.id,
            version_number=self.repository.next_version_number(resume.id),
            parent_version_id=parent.id if parent else None,
            source_type=source_type,
            source_file_id=source_file_id,
            schema_version=validated.schema_version,
            content_json=validated.model_dump(mode="json"),
            parsed_text=parsed_text,
            change_summary=change_summary,
            status="draft",
            created_by=self.user_id,
        )
        self.db.add(version)
        self.db.flush()
        if resume.active_version_id is None:
            resume.active_version_id = version.id
        self._audit("resume.version_created", resume.id, {"version_number": version.version_number})
        return _serialize(version, include_content=True)

    def finalize(self, resume_id: UUID, version_id: UUID) -> dict[str, object]:
        resume = self._resume_for_update(resume_id)
        version = self._version(resume.id, version_id)
        if version.status == "final":
            return _serialize(version, include_content=True)
        if version.status != "draft":
            raise ResumeConflict("Only a Draft Version can be finalized.")
        version.status = "final"
        version.finalized_at = utc_now()
        resume.active_version_id = version.id
        self._audit("resume.version_finalized", resume.id, {"version_number": version.version_number})
        return _serialize(version, include_content=True)

    def set_active(self, resume_id: UUID, version_id: UUID) -> dict[str, object]:
        resume = self._resume_for_update(resume_id)
        version = self._version(resume.id, version_id)
        resume.active_version_id = version.id
        self._audit("resume.active_version_set", resume.id, {"version_number": version.version_number})
        return _serialize(resume)

    def diff(self, resume_id: UUID, from_id: UUID, to_id: UUID) -> dict[str, object]:
        resume = self._resume(resume_id)
        before = self._version(resume.id, from_id)
        after = self._version(resume.id, to_id)
        return {
            "from_version_id": str(before.id),
            "to_version_id": str(after.id),
            "changes": _structured_diff(before.content_json, after.content_json),
        }

    def upload_file(self, filename: str, media_type: str, data: bytes) -> tuple[FileAsset, bool]:
        display, extension, digest = validate_resume_upload(
            filename, media_type, data, self.settings.max_stored_file_size_bytes
        )
        duplicate = self.repository.duplicate_file(self.user_id, digest)
        if duplicate:
            return duplicate, True
        storage_key, written_digest = self.storage.write(extension, data)
        asset = FileAsset(
            user_id=self.user_id,
            kind="resume",
            original_filename=display,
            storage_key=storage_key,
            media_type=media_type,
            size_bytes=len(data),
            sha256=written_digest,
        )
        self.db.add(asset)
        self.db.flush()
        self._audit("file.uploaded", asset.id, {"media_type": media_type, "size_bytes": len(data)})
        return asset, False

    def file(self, file_id: UUID) -> FileAsset:
        asset = self.repository.file(self.user_id, file_id)
        if asset is None:
            raise ResumeNotFound("File not found.")
        return asset

    def file_path(self, asset: FileAsset) -> Path:
        path = self.storage.path(asset.storage_key)
        if not path.is_file() or path.is_symlink():
            raise ResumeNotFound("Stored file is unavailable.")
        return path

    def delete_file(self, file_id: UUID) -> None:
        asset = self.file(file_id)
        if self.repository.file_is_referenced(asset.id):
            raise ResumeConflict("Referenced Resume files cannot be deleted.")
        asset.deleted_at = utc_now()
        self._audit("file.soft_deleted", asset.id)

    def import_resume(self, filename: str, media_type: str, data: bytes) -> dict[str, object]:
        # Parse before making any database or primary-resume changes. A failed
        # upload therefore cannot disturb the current primary resume.
        validate_resume_upload(filename, media_type, data, self.settings.max_stored_file_size_bytes)
        parsed_text = extract_resume_text_safely(data, media_type)
        asset, duplicate = self.upload_file(filename, media_type, data)
        title = Path(asset.original_filename).stem[:240] or "Imported Resume"
        resume = Resume(user_id=self.user_id, title=title, language="en", target_role="")
        self.db.add(resume)
        self.db.flush()
        content = conservative_resume_structure(parsed_text)
        version = self.create_version(
            resume.id,
            content,
            None,
            "Imported from an uploaded resume; manual review required.",
            parsed_text=parsed_text,
            source_type="import",
            source_file_id=asset.id,
        )
        self._make_primary(resume)
        self._audit("resume.imported", resume.id, {"duplicate_file": duplicate})
        return {
            "resume": _serialize(resume),
            "version": version,
            "file": _serialize(asset),
            "needs_review": True,
            "is_primary": True,
        }

    def _make_primary(self, resume: Resume, *, locked: list[Resume] | None = None) -> None:
        values = locked if locked is not None else self.repository.active_for_update(self.user_id)
        for value in values:
            value.is_primary = False
        # Flush the previous primary off before enabling the new one so the
        # partial unique index is respected on SQLite and PostgreSQL alike.
        self.db.flush()
        resume.is_primary = True
        self.db.flush()
        self._audit("resume.primary_set", resume.id)

    def _resume(self, resume_id: UUID) -> Resume:
        value = self.repository.resume(self.user_id, resume_id)
        if value is None:
            raise ResumeNotFound("Resume not found.")
        return value

    def _resume_for_update(self, resume_id: UUID) -> Resume:
        value = self.repository.resume_for_update(self.user_id, resume_id)
        if value is None:
            raise ResumeNotFound("Resume not found.")
        return value

    def _version(self, resume_id: UUID, version_id: UUID) -> ResumeVersion:
        value = self.repository.version(resume_id, version_id)
        if value is None:
            raise ResumeNotFound("Resume Version not found.")
        return value

    def _audit(self, event: str, resource_id: UUID, metadata: dict[str, object] | None = None) -> None:
        resource_type = "file_asset" if event.startswith("file.") else "resume"
        AuthRepository(self.db).audit(
            event,
            user_id=self.user_id,
            resource_type=resource_type,
            resource_id=str(resource_id),
            safe_metadata=metadata,
        )


def extract_resume_text_safely(data: bytes, media_type: str) -> str:
    try:
        if media_type == PDF_MEDIA_TYPE:
            reader = PdfReader(io.BytesIO(data), strict=True)
            if len(reader.pages) > 100:
                raise ValueError("PDF contains too many pages.")
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        elif media_type == DOCX_MEDIA_TYPE:
            document = Document(io.BytesIO(data))
            values = [paragraph.text for paragraph in document.paragraphs]
            values.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            text = "\n".join(values)
        elif media_type in {TEXT_MEDIA_TYPE, MARKDOWN_MEDIA_TYPE, "text/x-markdown"}:
            try:
                text = data.decode("utf-8", errors="strict")
            except UnicodeDecodeError:
                detected = chardet.detect(data).get("encoding") or "utf-8"
                try:
                    text = data.decode(detected, errors="replace")
                except LookupError:
                    text = data.decode("utf-8", errors="replace")
        else:
            raise ValueError("Unsupported resume media type.")
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("Resume document could not be parsed.") from exc
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        if media_type == PDF_MEDIA_TYPE:
            raise ValueError("No selectable text was found in this PDF.")
        raise ValueError("Resume does not contain extractable text.")
    if len(text) > 200_000:
        raise ValueError("Extracted resume text is too large.")
    return text


def conservative_resume_structure(text: str) -> dict[str, object]:
    section_names = {
        "experience": "experience",
        "work experience": "experience",
        "education": "education",
        "projects": "projects",
        "skills": "skills",
        "languages": "languages",
        "certifications": "certifications",
    }
    sections: list[dict[str, object]] = []
    current = {"type": "custom", "title": "Imported Content", "items": []}
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        normalized = line.lower().rstrip(":")
        if normalized in section_names:
            if current["items"]:
                sections.append(current)
            current = {"type": section_names[normalized], "title": line.rstrip(":"), "items": []}
            continue
        current["items"].append({"text": line[:2000], "verification_status": "needs_review"})
    if current["items"]:
        sections.append(current)
    return {"schema_version": 1, "header": {}, "summary": "", "sections": sections[:50]}


def _structured_diff(before: object, after: object, path: str = "$") -> list[dict[str, object]]:
    if before == after:
        return []
    if isinstance(before, dict) and isinstance(after, dict):
        changes: list[dict[str, object]] = []
        for key in sorted(set(before) | set(after)):
            changes.extend(_structured_diff(before.get(key), after.get(key), f"{path}.{key}"))
        return changes
    if isinstance(before, list) and isinstance(after, list):
        changes = []
        for index in range(max(len(before), len(after))):
            old = before[index] if index < len(before) else None
            new = after[index] if index < len(after) else None
            changes.extend(_structured_diff(old, new, f"{path}[{index}]"))
        return changes
    return [{"path": path, "before": before, "after": after}]
