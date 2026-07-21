import io
import os
import tempfile
import unittest
import zipfile
from datetime import timedelta
from pathlib import Path
from unittest.mock import patch
from uuid import UUID

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from reportlab.pdfgen import canvas
from sqlalchemy import select

from app.api.routers import auth, profile, resumes
from app.auth.middleware import V2SecurityMiddleware
from app.auth.service import AuthService
from app.core.config import V2ConfigError, load_v2_settings
from app.core.security import hash_password, verify_password
from app.db.base import Base
from app.db.engine import build_engine
from app.db.models import ResumeVersion, UserSession, utc_now
from app.db.session import session_factory
from app.resumes.service import ResumeService


def docx_bytes(text="Platform Engineer"):
    document = Document()
    document.add_heading("Experience", 1)
    document.add_paragraph(text)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def malformed_docx_bytes():
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>',
        )
        archive.writestr(
            "word/document.xml",
            '<?xml version="1.0"?><w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:p><w:r><w:t>Incomplete package</w:t></w:r></w:p></w:body></w:document>',
        )
    return output.getvalue()


def pdf_bytes(text="Platform Engineer"):
    output = io.BytesIO()
    document = canvas.Canvas(output)
    if text:
        document.drawString(72, 720, text)
    document.showPage()
    document.save()
    return output.getvalue()


class V2FoundationTest(unittest.TestCase):
    def setUp(self):
        self.temporary = tempfile.TemporaryDirectory()
        root = Path(self.temporary.name)
        self.environment = patch.dict(
            os.environ,
            {
                "APP_ENV": "test",
                "AUTH_ENABLED": "true",
                "TEST_DATABASE_URL": f"sqlite+pysqlite:///{root / 'phase1-test.db'}",
                "FILE_STORAGE_ROOT": str(root / "files"),
                "SESSION_COOKIE_SECURE": "false",
                "AUTH_TRUSTED_ORIGINS": "http://testserver",
                "AUTH_FINGERPRINT_KEY": "TEST_ONLY_FINGERPRINT_KEY_32_BYTES_LONG",
                "AUTH_MAX_FAILED_ATTEMPTS": "3",
                "AUTH_LOCKOUT_MINUTES": "5",
            },
        )
        self.environment.start()
        build_engine.cache_clear()
        self.settings = load_v2_settings()
        self.engine = build_engine(self.settings.database_url)
        Base.metadata.create_all(self.engine)
        db = session_factory(self.settings.database_url)()
        try:
            service = AuthService(db, self.settings)
            self.admin = service.create_user("admin@example.com", "correct horse battery staple", "Admin", "admin")
            self.other = service.create_user("other@example.com", "another correct passphrase", "Other", "user")
            self.admin_id = self.admin.id
            self.other_id = self.other.id
            db.commit()
        finally:
            db.close()
        app = FastAPI()
        app.include_router(auth.router)
        app.include_router(profile.router)
        app.include_router(resumes.router)
        app.add_middleware(V2SecurityMiddleware, settings=self.settings)
        self.app = app
        self.client = TestClient(app)

    def tearDown(self):
        self.client.close()
        self.engine.dispose()
        build_engine.cache_clear()
        self.environment.stop()
        self.temporary.cleanup()

    def login(self, client=None, email="admin@example.com", password="correct horse battery staple"):
        client = client or self.client
        response = client.post("/api/auth/login", json={"email": email, "password": password})
        self.assertEqual(response.status_code, 200, response.text)
        return response.json()["csrf_token"]

    def unsafe_headers(self, csrf):
        return {"Origin": "http://testserver", "X-CSRF-Token": csrf}

    def test_password_hash_is_argon2_and_not_plaintext(self):
        encoded = hash_password("long and memorable passphrase")
        self.assertNotEqual(encoded, "long and memorable passphrase")
        self.assertTrue(encoded.startswith("$argon2"))
        self.assertTrue(verify_password("long and memorable passphrase", encoded))
        self.assertFalse(verify_password("wrong passphrase", encoded))

    def test_production_rejects_insecure_cookie(self):
        with patch.dict(
            os.environ,
            {
                "APP_ENV": "production",
                "DATABASE_URL": "postgresql+psycopg://user:fake@db/app",
                "SESSION_COOKIE_SECURE": "false",
                "AUTH_TRUSTED_ORIGINS": "https://example.test",
                "AUTH_FINGERPRINT_KEY": "x" * 32,
            },
        ):
            with self.assertRaises(V2ConfigError):
                load_v2_settings()

    def test_login_error_is_generic_for_unknown_and_wrong_password(self):
        unknown = self.client.post("/api/auth/login", json={"email": "missing@example.com", "password": "not the password"})
        wrong = self.client.post("/api/auth/login", json={"email": "admin@example.com", "password": "not the password"})
        self.assertEqual(unknown.status_code, 401)
        self.assertEqual(unknown.json(), wrong.json())

    def test_inactive_user_cannot_login(self):
        db = session_factory(self.settings.database_url)()
        try:
            user = db.get(type(self.other), self.other_id)
            user.is_active = False
            db.commit()
        finally:
            db.close()
        response = self.client.post("/api/auth/login", json={"email": "other@example.com", "password": "another correct passphrase"})
        self.assertEqual(response.status_code, 401)

    def test_login_cookie_is_httponly_lax_and_token_is_hashed(self):
        response = self.client.post("/api/auth/login", json={"email": "admin@example.com", "password": "correct horse battery staple"})
        cookie = response.headers["set-cookie"]
        self.assertIn("HttpOnly", cookie)
        self.assertIn("SameSite=lax", cookie)
        raw = self.client.cookies.get("pja_session")
        db = session_factory(self.settings.database_url)()
        try:
            stored = db.scalar(select(UserSession).order_by(UserSession.created_at.desc()))
            self.assertNotEqual(raw, stored.token_hash)
            self.assertNotIn(raw, stored.token_hash)
        finally:
            db.close()

    def test_normal_and_remembered_session_ttls_are_bounded(self):
        normal_response = self.client.post(
            "/api/auth/login",
            json={"email": "admin@example.com", "password": "correct horse battery staple", "remember_me": False},
        )
        self.assertNotIn("Max-Age", normal_response.headers["set-cookie"])
        db = session_factory(self.settings.database_url)()
        try:
            normal = db.scalar(select(UserSession).order_by(UserSession.created_at.desc()))
            normal_idle = normal.idle_expires_at - normal.created_at
            normal_absolute = normal.absolute_expires_at - normal.created_at
            self.assertAlmostEqual(normal_idle.total_seconds(), 30 * 60, delta=5)
            self.assertAlmostEqual(normal_absolute.total_seconds(), 24 * 3600, delta=5)
        finally:
            db.close()

        remembered_response = self.client.post(
            "/api/auth/login",
            json={"email": "admin@example.com", "password": "correct horse battery staple", "remember_me": True},
        )
        cookie = remembered_response.headers["set-cookie"]
        self.assertIn("Max-Age=2592000", cookie)
        db = session_factory(self.settings.database_url)()
        try:
            remembered = db.scalar(select(UserSession).order_by(UserSession.created_at.desc()))
            lifetime = remembered.absolute_expires_at - remembered.created_at
            self.assertAlmostEqual(lifetime.total_seconds(), 30 * 86400, delta=5)
            self.assertEqual(remembered.idle_expires_at, remembered.absolute_expires_at)
        finally:
            db.close()

    def test_login_rotates_previous_session(self):
        self.login()
        first_token = self.client.cookies.get("pja_session")
        self.login()
        second_token = self.client.cookies.get("pja_session")
        self.assertNotEqual(first_token, second_token)
        db = session_factory(self.settings.database_url)()
        try:
            revoked = db.scalar(select(UserSession).where(UserSession.token_hash != "").order_by(UserSession.created_at.asc()))
            self.assertIsNotNone(revoked.revoked_at)
            self.assertEqual(revoked.revoke_reason, "login_rotation")
        finally:
            db.close()

    def test_password_change_revokes_old_sessions_and_rotates_current(self):
        csrf = self.login()
        old_token = self.client.cookies.get("pja_session")
        response = self.client.post(
            "/api/auth/change-password",
            json={"current_password": "correct horse battery staple", "new_password": "new correct horse battery staple"},
            headers=self.unsafe_headers(csrf),
        )
        self.assertEqual(response.status_code, 200, response.text)
        self.assertNotEqual(old_token, self.client.cookies.get("pja_session"))
        db = session_factory(self.settings.database_url)()
        try:
            sessions = list(db.scalars(select(UserSession).order_by(UserSession.created_at.asc())))
            self.assertTrue(all(item.revoked_at is not None for item in sessions[:-1]))
            self.assertEqual(sessions[-1].revoked_at, None)
        finally:
            db.close()

    def test_absolute_expiry_rejects_a_remembered_session(self):
        self.client.post(
            "/api/auth/login",
            json={"email": "admin@example.com", "password": "correct horse battery staple", "remember_me": True},
        )
        db = session_factory(self.settings.database_url)()
        try:
            value = db.scalar(select(UserSession).order_by(UserSession.created_at.desc()))
            value.absolute_expires_at = utc_now() - timedelta(seconds=1)
            db.commit()
        finally:
            db.close()
        self.assertEqual(self.client.get("/api/profile").status_code, 401)

    def test_session_bootstrap_rotates_and_returns_csrf_only(self):
        login_csrf = self.login()
        payload = self.client.get("/api/auth/session").json()
        self.assertTrue(payload["authenticated"])
        self.assertNotEqual(payload["csrf_token"], login_csrf)
        self.assertNotIn("token", payload["user"])
        self.assertNotIn("password_hash", payload["user"])

    def test_protected_api_requires_authentication(self):
        self.assertEqual(self.client.get("/api/profile").status_code, 401)

    def test_csrf_and_origin_are_required_for_unsafe_methods(self):
        csrf = self.login()
        profile_payload = {"revision": 1, "headline": "Engineer"}
        self.assertEqual(self.client.put("/api/profile", json=profile_payload).status_code, 403)
        self.assertEqual(self.client.put("/api/profile", json=profile_payload, headers={"Origin": "http://evil.test", "X-CSRF-Token": csrf}).status_code, 403)
        self.assertEqual(self.client.put("/api/profile?csrf_token=" + csrf, json=profile_payload, headers={"Origin": "http://testserver"}).status_code, 403)

    def test_login_rate_limit_is_database_backed(self):
        for _ in range(3):
            self.client.post("/api/auth/login", json={"email": "admin@example.com", "password": "wrong"})
        response = self.client.post("/api/auth/login", json={"email": "admin@example.com", "password": "correct horse battery staple"})
        self.assertEqual(response.status_code, 429)

    def test_logout_revokes_server_session(self):
        csrf = self.login()
        response = self.client.post("/api/auth/logout", headers=self.unsafe_headers(csrf))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(self.client.get("/api/profile").status_code, 401)

    def test_expired_session_is_rejected(self):
        self.login()
        db = session_factory(self.settings.database_url)()
        try:
            value = db.scalar(select(UserSession))
            value.idle_expires_at = utc_now() - timedelta(seconds=1)
            db.commit()
        finally:
            db.close()
        self.assertEqual(self.client.get("/api/profile").status_code, 401)

    def test_profile_crud_revision_and_stale_update(self):
        csrf = self.login()
        profile_value = self.client.get("/api/profile").json()
        updated = self.client.put(
            "/api/profile",
            json={"revision": profile_value["revision"], "headline": "Platform Engineer", "professional_summary": "Builds secure systems", "current_location": "Remote"},
            headers=self.unsafe_headers(csrf),
        )
        self.assertEqual(updated.status_code, 200, updated.text)
        current_revision = updated.json()["revision"]
        created = self.client.post(
            "/api/profile/experiences",
            json={"company": "Example", "role_title": "Engineer", "is_current": True, "verification_status": "confirmed"},
            headers={**self.unsafe_headers(csrf), "If-Match": str(current_revision)},
        )
        self.assertEqual(created.status_code, 201, created.text)
        stale = self.client.post(
            "/api/profile/skills",
            json={"name": "Python"},
            headers={**self.unsafe_headers(csrf), "If-Match": str(current_revision)},
        )
        self.assertEqual(stale.status_code, 409)
        self.assertTrue(self.client.get("/api/profile/revisions").json())

    def test_profile_revision_restore_creates_new_revision(self):
        csrf = self.login()
        first = self.client.get("/api/profile").json()
        second = self.client.put("/api/profile", json={"revision": first["revision"], "headline": "First"}, headers=self.unsafe_headers(csrf)).json()
        third = self.client.put("/api/profile", json={"revision": second["revision"], "headline": "Second"}, headers=self.unsafe_headers(csrf)).json()
        restored = self.client.post(f"/api/profile/revisions/{second['revision']}/restore", headers={**self.unsafe_headers(csrf), "If-Match": str(third["revision"])})
        self.assertEqual(restored.status_code, 200, restored.text)
        self.assertEqual(restored.json()["headline"], "First")
        self.assertGreater(restored.json()["revision"], third["revision"])

    def test_resume_version_finalize_diff_and_immutability(self):
        csrf = self.login()
        resume = self.client.post("/api/resumes", json={"title": "Primary Resume"}, headers=self.unsafe_headers(csrf)).json()
        first = self.client.post(f"/api/resumes/{resume['id']}/versions", json={"content": {"schema_version": 1, "header": {}, "summary": "One", "sections": []}}, headers=self.unsafe_headers(csrf)).json()
        second = self.client.post(f"/api/resumes/{resume['id']}/versions", json={"parent_version_id": first["id"], "content": {"schema_version": 1, "header": {}, "summary": "Two", "sections": []}}, headers=self.unsafe_headers(csrf)).json()
        difference = self.client.get(f"/api/resumes/{resume['id']}/diff?from_version_id={first['id']}&to_version_id={second['id']}")
        self.assertTrue(difference.json()["changes"])
        finalized = self.client.post(f"/api/resumes/{resume['id']}/versions/{second['id']}/finalize", headers=self.unsafe_headers(csrf))
        self.assertEqual(finalized.json()["status"], "final")
        db = session_factory(self.settings.database_url)()
        try:
            version = db.get(ResumeVersion, UUID(second["id"]))
            version.content_json = {"schema_version": 1, "header": {}, "summary": "Mutated", "sections": []}
            with self.assertRaises(ValueError):
                db.commit()
            db.rollback()
        finally:
            db.close()

    def test_resume_version_lineage_metadata_is_immutable(self):
        csrf = self.login()
        resume = self.client.post(
            "/api/resumes",
            json={"title": "Metadata Resume"},
            headers=self.unsafe_headers(csrf),
        ).json()
        version_value = self.client.post(
            f"/api/resumes/{resume['id']}/versions",
            json={"content": {"schema_version": 1, "header": {}, "summary": "One", "sections": []}},
            headers=self.unsafe_headers(csrf),
        ).json()
        db = session_factory(self.settings.database_url)()
        try:
            version = db.get(ResumeVersion, UUID(version_value["id"]))
            version.change_summary = "Mutated lineage"
            with self.assertRaises(ValueError):
                db.commit()
            db.rollback()
        finally:
            db.close()

    def test_cross_user_resume_and_file_idor_are_not_found(self):
        other_client = TestClient(self.app)
        try:
            other_csrf = self.login(other_client, "other@example.com", "another correct passphrase")
            resume = other_client.post("/api/resumes", json={"title": "Other Resume"}, headers=self.unsafe_headers(other_csrf)).json()
            upload = other_client.post("/api/files/resume", files={"file": ("resume.docx", docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=self.unsafe_headers(other_csrf)).json()
            self.login()
            self.assertEqual(self.client.get(f"/api/resumes/{resume['id']}").status_code, 404)
            self.assertEqual(self.client.get(f"/api/files/{upload['file']['id']}/metadata").status_code, 404)
            self.assertIsNone(self.client.get("/api/resumes/primary").json())
        finally:
            other_client.close()

    def test_docx_upload_download_duplicate_and_reference_protection(self):
        csrf = self.login()
        content = docx_bytes()
        first = self.client.post("/api/files/resume", files={"file": ("resume.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=self.unsafe_headers(csrf))
        second = self.client.post("/api/files/resume", files={"file": ("resume.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=self.unsafe_headers(csrf))
        self.assertFalse(first.json()["duplicate"])
        self.assertTrue(second.json()["duplicate"])
        file_id = first.json()["file"]["id"]
        download = self.client.get(f"/api/files/{file_id}/download")
        self.assertEqual(download.content, content)
        self.assertEqual(download.headers["x-content-type-options"], "nosniff")
        imported = self.client.post("/api/resumes/import", files={"file": ("resume.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=self.unsafe_headers(csrf))
        self.assertEqual(imported.status_code, 201, imported.text)
        self.assertEqual(self.client.delete(f"/api/files/{file_id}", headers=self.unsafe_headers(csrf)).status_code, 409)

    def test_upload_rejects_mime_spoof(self):
        csrf = self.login()
        response = self.client.post("/api/files/resume", files={"file": ("resume.pdf", b"not a pdf", "application/pdf")}, headers=self.unsafe_headers(csrf))
        self.assertEqual(response.status_code, 400)

    def test_import_rejects_malformed_docx_package_without_server_error(self):
        csrf = self.login()
        response = self.client.post(
            "/api/resumes/import",
            files={
                "file": (
                    "resume.docx",
                    malformed_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            headers=self.unsafe_headers(csrf),
        )
        self.assertEqual(response.status_code, 400, response.text)

    def test_pdf_upload_creates_primary_resume(self):
        csrf = self.login()
        response = self.client.post(
            "/api/resumes/upload",
            files={"file": ("platform.pdf", pdf_bytes("Python Platform Engineer"), "application/pdf")},
            headers=self.unsafe_headers(csrf),
        )
        self.assertEqual(response.status_code, 201, response.text)
        value = response.json()
        self.assertTrue(value["resume"]["is_primary"])
        self.assertIn("Python Platform Engineer", value["version"]["extracted_text"])
        self.assertEqual(value["file"]["original_filename"], "platform.pdf")
        self.assertTrue(value["file"]["content_hash"])

    def test_txt_upload_creates_resume_and_preserves_text(self):
        csrf = self.login()
        response = self.client.post(
            "/api/resumes/upload",
            files={"file": ("resume.txt", "Skills\r\n• Python\r\nFastAPI".encode(), "text/plain")},
            headers=self.unsafe_headers(csrf),
        )
        self.assertEqual(response.status_code, 201, response.text)
        self.assertIn("• Python\nFastAPI", response.json()["version"]["extracted_text"])

    def test_markdown_upload_is_supported(self):
        csrf = self.login()
        response = self.client.post(
            "/api/resumes/upload",
            files={"file": ("resume.md", b"# Skills\n- Python", "text/markdown")},
            headers=self.unsafe_headers(csrf),
        )
        self.assertEqual(response.status_code, 201, response.text)

    def test_unsupported_upload_type_is_rejected(self):
        csrf = self.login()
        response = self.client.post(
            "/api/resumes/upload",
            files={"file": ("resume.doc", b"legacy", "application/msword")},
            headers=self.unsafe_headers(csrf),
        )
        self.assertEqual(response.status_code, 400)

    def test_oversized_upload_is_rejected(self):
        csrf = self.login()
        with patch.dict(os.environ, {"MAX_STORED_FILE_SIZE_MB": "1"}, clear=False):
            response = self.client.post(
                "/api/resumes/upload",
                files={"file": ("resume.txt", b"x" * (1024 * 1024 + 1), "text/plain")},
                headers=self.unsafe_headers(csrf),
            )
        self.assertEqual(response.status_code, 400)

    def test_pdf_without_selectable_text_has_clear_error(self):
        csrf = self.login()
        response = self.client.post(
            "/api/resumes/upload",
            files={"file": ("scan.pdf", pdf_bytes(""), "application/pdf")},
            headers=self.unsafe_headers(csrf),
        )
        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.json()["detail"], "No selectable text was found in this PDF.")

    def test_latest_upload_is_only_primary_and_primary_api_returns_it(self):
        csrf = self.login()
        first = self.client.post(
            "/api/resumes/upload", files={"file": ("first.txt", b"Python", "text/plain")}, headers=self.unsafe_headers(csrf)
        ).json()
        second = self.client.post(
            "/api/resumes/upload", files={"file": ("second.txt", b"FastAPI", "text/plain")}, headers=self.unsafe_headers(csrf)
        ).json()
        resumes_value = self.client.get("/api/resumes").json()
        self.assertEqual([item["id"] for item in resumes_value if item["is_primary"]], [second["resume"]["id"]])
        self.assertEqual(self.client.get("/api/resumes/primary").json()["id"], second["resume"]["id"])
        self.assertFalse(next(item for item in resumes_value if item["id"] == first["resume"]["id"])["is_primary"])

    def test_failed_upload_does_not_change_primary(self):
        csrf = self.login()
        primary = self.client.post(
            "/api/resumes/upload", files={"file": ("valid.txt", b"Python", "text/plain")}, headers=self.unsafe_headers(csrf)
        ).json()["resume"]
        failed = self.client.post(
            "/api/resumes/upload", files={"file": ("scan.pdf", pdf_bytes(""), "application/pdf")}, headers=self.unsafe_headers(csrf)
        )
        self.assertEqual(failed.status_code, 400)
        self.assertEqual(self.client.get("/api/resumes/primary").json()["id"], primary["id"])

    def test_deleting_primary_selects_latest_remaining_resume(self):
        csrf = self.login()
        first = self.client.post(
            "/api/resumes/upload", files={"file": ("first.txt", b"Python", "text/plain")}, headers=self.unsafe_headers(csrf)
        ).json()["resume"]
        second = self.client.post(
            "/api/resumes/upload", files={"file": ("second.txt", b"FastAPI", "text/plain")}, headers=self.unsafe_headers(csrf)
        ).json()["resume"]
        deleted = self.client.delete(f"/api/resumes/{second['id']}", headers=self.unsafe_headers(csrf))
        self.assertEqual(deleted.status_code, 200)
        self.assertEqual(self.client.get("/api/resumes/primary").json()["id"], first["id"])

    def test_duplicate_upload_keeps_one_primary_and_valid_versions(self):
        csrf = self.login()
        content = b"Python FastAPI"
        first = self.client.post(
            "/api/resumes/upload", files={"file": ("same.txt", content, "text/plain")}, headers=self.unsafe_headers(csrf)
        ).json()
        second = self.client.post(
            "/api/resumes/upload", files={"file": ("same.txt", content, "text/plain")}, headers=self.unsafe_headers(csrf)
        ).json()
        self.assertEqual(first["file"]["id"], second["file"]["id"])
        self.assertNotEqual(first["resume"]["id"], second["resume"]["id"])
        self.assertEqual(len([item for item in self.client.get("/api/resumes").json() if item["is_primary"]]), 1)


if __name__ == "__main__":
    unittest.main()
